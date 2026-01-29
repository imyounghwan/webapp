"""
KRDS Word 문서에서 이미지 추출 및 분석
- Word 문서 내부 이미지 추출
- 잘한 점/못한 점 구분
- 이미지별 메타데이터 저장
"""

import os
import json
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

class KRDSImageExtractor:
    def __init__(self, docx_dir: str, output_dir: str):
        self.docx_dir = Path(docx_dir)
        self.output_dir = Path(output_dir)
        self.images_dir = self.output_dir / "images"
        self.images_dir.mkdir(parents=True, exist_ok=True)
        
    def extract_images_from_docx(self, docx_path: Path) -> list:
        """Word 문서에서 모든 이미지 추출"""
        doc = Document(docx_path)
        agency_name = self._extract_agency_name(docx_path.name)
        images = []
        
        # 이미지 카운터
        image_counter = 1
        
        # 문서 내 모든 관계(relationships) 확인
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    # 이미지 데이터 추출
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # 확장자 추출
                    content_type = image_part.content_type
                    ext = self._get_extension(content_type)
                    
                    # 파일명 생성
                    filename = f"{agency_name}_{image_counter:02d}{ext}"
                    image_path = self.images_dir / filename
                    
                    # 이미지 저장
                    with open(image_path, 'wb') as f:
                        f.write(image_data)
                    
                    images.append({
                        "agency": agency_name,
                        "filename": filename,
                        "path": str(image_path),
                        "size": len(image_data),
                        "type": content_type,
                        "index": image_counter
                    })
                    
                    image_counter += 1
                    
                except Exception as e:
                    print(f"⚠️  이미지 추출 실패 ({docx_path.name}): {e}")
        
        return images
    
    def extract_text_context(self, docx_path: Path) -> dict:
        """Word 문서에서 텍스트 컨텍스트 추출 (잘한 점/못한 점 구분용)"""
        doc = Document(docx_path)
        agency_name = self._extract_agency_name(docx_path.name)
        
        context = {
            "agency": agency_name,
            "good_points": [],
            "bad_points": [],
            "sections": []
        }
        
        current_section = None
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 섹션 구분
            if any(keyword in text for keyword in ["잘한 점", "우수 사례", "긍정", "좋은"]):
                current_section = "good"
                context["sections"].append({"type": "good", "title": text})
            elif any(keyword in text for keyword in ["못한 점", "개선", "문제", "미흡", "취약"]):
                current_section = "bad"
                context["sections"].append({"type": "bad", "title": text})
            else:
                # 현재 섹션에 텍스트 추가
                if current_section == "good":
                    context["good_points"].append(text)
                elif current_section == "bad":
                    context["bad_points"].append(text)
        
        return context
    
    def process_all_documents(self) -> dict:
        """모든 KRDS Word 문서 처리"""
        all_results = {
            "agencies": [],
            "total_images": 0,
            "images_by_agency": {}
        }
        
        # KRDS Word 문서 목록
        docx_files = list(self.docx_dir.glob("*.docx"))
        
        print(f"📂 KRDS Word 문서 발견: {len(docx_files)}개")
        
        for docx_file in sorted(docx_files):
            print(f"\n🔍 처리 중: {docx_file.name}")
            
            try:
                # 이미지 추출
                images = self.extract_images_from_docx(docx_file)
                
                # 텍스트 컨텍스트 추출
                context = self.extract_text_context(docx_file)
                
                agency_name = self._extract_agency_name(docx_file.name)
                
                agency_result = {
                    "agency": agency_name,
                    "docx_path": str(docx_file),
                    "images": images,
                    "image_count": len(images),
                    "context": context
                }
                
                all_results["agencies"].append(agency_result)
                all_results["total_images"] += len(images)
                all_results["images_by_agency"][agency_name] = len(images)
                
                print(f"✅ {agency_name}: {len(images)}개 이미지 추출")
                
            except Exception as e:
                print(f"❌ 처리 실패 ({docx_file.name}): {e}")
        
        return all_results
    
    def save_results(self, results: dict):
        """결과 JSON 저장"""
        output_file = self.output_dir / "krds_images_metadata.json"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        
        print(f"\n💾 메타데이터 저장: {output_file}")
        
        # 요약 통계
        print(f"\n📊 추출 요약:")
        print(f"   총 기관: {len(results['agencies'])}개")
        print(f"   총 이미지: {results['total_images']}개")
        print(f"   평균 이미지/기관: {results['total_images'] / len(results['agencies']):.1f}개")
    
    def _extract_agency_name(self, filename: str) -> str:
        """파일명에서 기관명 추출"""
        # 예: "해양수산부 - 연안포털 - 웹사이트 품질관리 수준진단 보고서.docx"
        #     → "해양수산부 - 연안포털"
        parts = filename.replace(".docx", "").split(" - ")
        if len(parts) >= 2:
            return f"{parts[0]} - {parts[1]}"
        return parts[0]
    
    def _get_extension(self, content_type: str) -> str:
        """MIME 타입에서 확장자 추출"""
        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
        }
        return ext_map.get(content_type, ".png")


def main():
    # 디렉토리 설정
    docx_dir = "/home/user/uploaded_files"
    output_dir = "/home/user/webapp/analysis/krds_images"
    
    print("🚀 KRDS 이미지 추출 시작...\n")
    
    # 추출기 초기화
    extractor = KRDSImageExtractor(docx_dir, output_dir)
    
    # 모든 문서 처리
    results = extractor.process_all_documents()
    
    # 결과 저장
    extractor.save_results(results)
    
    print("\n✨ 완료!")


if __name__ == "__main__":
    main()
