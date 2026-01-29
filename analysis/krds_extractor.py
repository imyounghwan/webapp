#!/usr/bin/env python3
"""
KRDS 편의성 점수 추출기
Word 문서에서 '웹 편의성' 점수를 추출합니다.
"""

import os
import json
import re
from docx import Document
from pathlib import Path

def extract_krds_score(docx_path):
    """Word 문서에서 KRDS 편의성 점수 추출"""
    try:
        doc = Document(docx_path)
        
        # 파일명에서 기관명 추출
        filename = Path(docx_path).name
        # "기관명 - 사이트명 - 웹사이트 품질관리 수준진단 보고서.docx" 형식
        parts = filename.replace(' - 웹사이트 품질관리 수준진단 보고서.docx', '').split(' - ')
        
        if len(parts) >= 2:
            department = parts[0].strip()
            site_name = parts[1].strip()
        else:
            department = parts[0].strip() if parts else filename
            site_name = "대표누리집"
        
        # 표에서 '웹 편의성' 점수 찾기
        convenience_score = None
        
        for table in doc.tables:
            header_row = None
            convenience_col_idx = None
            
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                
                # 헤더 행 찾기 ("웹 편의성" 열 위치 확인)
                if '웹 편의성' in cells:
                    header_row = row_idx
                    convenience_col_idx = cells.index('웹 편의성')
                    continue
                
                # "진단 결과" 행에서 점수 추출
                if convenience_col_idx is not None and ('진단 결과' in cells[0] or '진단결과' in cells[0]):
                    if convenience_col_idx < len(cells):
                        score_text = cells[convenience_col_idx]
                        match = re.search(r'(\d+\.?\d*)', score_text)
                        if match:
                            convenience_score = float(match.group(1))
                            break
            
            if convenience_score is not None:
                break
        
        # 추가 검색: 본문에서 "웹 편의성" 키워드와 점수 찾기
        if convenience_score is None:
            for para in doc.paragraphs:
                text = para.text
                if '웹 편의성' in text or '웹편의성' in text:
                    # "웹 편의성: 92점" 또는 "웹 편의성 87.5" 같은 패턴
                    match = re.search(r'웹\s*편의성[:\s]*(\d+\.?\d*)', text)
                    if match:
                        convenience_score = float(match.group(1))
                        break
        
        return {
            'department': department,
            'site_name': site_name,
            'full_name': f"{department} - {site_name}",
            'krds_convenience': convenience_score,
            'source_file': filename
        }
        
    except Exception as e:
        print(f"❌ 오류 발생 ({docx_path}): {e}")
        return None

def main():
    """메인 실행 함수"""
    
    # KRDS Word 파일 디렉토리
    docx_dir = Path('/home/user/uploaded_files')
    output_dir = Path('/home/user/webapp/analysis/krds_data')
    output_dir.mkdir(parents=True, exist_ok=True)
    
    # 모든 Word 파일 처리
    krds_scores = []
    
    docx_files = list(docx_dir.glob('*.docx'))
    print(f"📄 총 {len(docx_files)}개의 Word 파일 발견\n")
    
    for docx_file in sorted(docx_files):
        print(f"📖 처리 중: {docx_file.name}")
        result = extract_krds_score(docx_file)
        
        if result and result['krds_convenience'] is not None:
            krds_scores.append(result)
            print(f"   ✅ {result['full_name']}: {result['krds_convenience']}점\n")
        else:
            print(f"   ⚠️ 편의성 점수를 찾을 수 없음\n")
    
    # JSON 저장
    output_file = output_dir / 'krds_convenience_scores.json'
    with open(output_file, 'w', encoding='utf-8') as f:
        json.dump(krds_scores, f, ensure_ascii=False, indent=2)
    
    print(f"\n✅ 총 {len(krds_scores)}개 기관의 KRDS 편의성 점수 추출 완료")
    print(f"📁 저장 위치: {output_file}")
    
    # 요약 통계
    if krds_scores:
        scores = [s['krds_convenience'] for s in krds_scores]
        print(f"\n📊 통계:")
        print(f"   최고점: {max(scores)}점")
        print(f"   최저점: {min(scores)}점")
        print(f"   평균: {sum(scores)/len(scores):.1f}점")
        
        # TOP 5
        print(f"\n🏆 최고점 TOP 5:")
        for item in sorted(krds_scores, key=lambda x: x['krds_convenience'], reverse=True)[:5]:
            print(f"   {item['full_name']}: {item['krds_convenience']}점")
        
        # BOTTOM 5
        print(f"\n⚠️ 최저점 BOTTOM 5:")
        for item in sorted(krds_scores, key=lambda x: x['krds_convenience'])[:5]:
            print(f"   {item['full_name']}: {item['krds_convenience']}점")

if __name__ == '__main__':
    main()
